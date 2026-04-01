#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GENERATEUR DE RAPPORT INTERACTIF
- Watermark (image de fond pleine page) sur chaque page
- Couverture fidelite au template (image intro + texte)
- Claude AI pour generation automatique des paragraphes
- Systeme de templates extensible via templates.json

INSTALLATION :
    pip install python-docx Pillow lxml

LANCER :
    PowerShell  :  $env:ANTHROPIC_API_KEY="sk-ant-..."
    CMD Windows :  set ANTHROPIC_API_KEY=sk-ant-...
    python report_tool.py

AJOUTER UN TEMPLATE :
    1. Mettre 3 images dans templates/  (monid_intro.png  monid_outro.png  monid_page.png)
    2. Ajouter un bloc dans templates.json
    3. Relancer le script
"""

import os, sys, json, textwrap, copy
import urllib.request, urllib.error
from pathlib import Path
from datetime import datetime

# ── dependances ───────────────────────────────────────────────
def check_deps():
    missing = []
    for pkg, imp in [("python-docx","docx"), ("Pillow","PIL"), ("lxml","lxml")]:
        try: __import__(imp)
        except ImportError: missing.append(pkg)
    if missing:
        print("[ERREUR] pip install " + " ".join(missing))
        sys.exit(1)
check_deps()

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

# ── chemins ───────────────────────────────────────────────────
SCRIPT_DIR     = Path(__file__).parent.resolve()
TEMPLATES_DIR  = SCRIPT_DIR / "templates"
TEMPLATES_JSON = SCRIPT_DIR / "templates.json"
OUTPUT_DIR     = SCRIPT_DIR / "rapports_generes"

# ── API Claude ────────────────────────────────────────────────
API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
API_URL = "https://api.anthropic.com/v1/messages"
MODEL   = "claude-opus-4-5"

# ── CONFIGURATION TYPOGRAPHIE ─────────────────────────────
FONT_NAME         = "Montserrat"   # police globale
FONT_SIZE_TITLE   = 33             # titres de section
FONT_SIZE_BODY    = 20             # paragraphes
FONT_SIZE_CAPTION = 15             # légendes figures
FONT_SIZE_COVER   = 53             # titre couverture
FONT_SIZE_REPORT  = 80             # mot REPORT
FONT_SIZE_SUBTITLE = 24             # sous-titres (ex: nom des étapes)


# =============================================================
#  UI HELPERS
# =============================================================

def banner():
    print()
    print("=" * 58)
    print("|" + " GENERATEUR DE RAPPORT INTERACTIF ".center(56) + "|")
    print("|" + " Templates personnalises + Claude AI ".center(56) + "|")
    print("=" * 58); print()

def section_title(t):
    print(); print("-" * 58)
    print("  " + t); print("-" * 58)

def ask(prompt, default=""):
    sfx = (" [" + default + "]") if default else ""
    try:   val = input("  >> " + prompt + sfx + " : ").strip()
    except EOFError: return default
    return val if val else default

def ask_yesno(prompt, default=True):
    h = "O/n" if default else "o/N"
    try:   val = input("  >> " + prompt + " (" + h + ") : ").strip().lower()
    except EOFError: return default
    if not val: return default
    return val in ("o", "oui", "y", "yes")

def progress(msg): print("  [...] " + msg, end="", flush=True)
def done_ok():     print(" OK")
def info(msg):     print("  (i) " + msg)
def warn(msg):     print("  /!\\ " + msg)

def wrap_print(text, width=52):
    for line in textwrap.wrap(str(text), width):
        print("      " + line)

def collect_manual(label):
    """Saisie manuelle multi-lignes, terminee par une ligne vide."""
    print("  Saisir " + label + " (ligne vide pour terminer) :")
    lines = []
    while True:
        try:   line = input("  > ")
        except EOFError: break
        if not line: break
        lines.append(line)
    return " ".join(lines) if lines else "(vide)"


# =============================================================
#  CHARGEMENT TEMPLATES
# =============================================================

def load_templates():
    if not TEMPLATES_JSON.exists():
        warn("templates.json introuvable a cote du script."); sys.exit(1)
    with open(TEMPLATES_JSON, encoding="utf-8") as f:
        data = json.load(f)
    valid = []
    for tpl in data.get("templates", []):
        intro = TEMPLATES_DIR / tpl["intro_bg"]
        outro = TEMPLATES_DIR / tpl["outro_bg"]
        page  = TEMPLATES_DIR / tpl["page_bg"]
        if intro.exists() and outro.exists() and page.exists():
            tpl["_intro"]     = intro
            tpl["_outro"]     = outro
            tpl["_page"]      = page
            tpl["_accent"]    = RGBColor(*tpl["accent_color"])
            tpl["_secondary"] = RGBColor(*tpl["secondary_color"])
            tpl["_dark"]      = RGBColor(*tpl["text_dark"])
            tpl["_subtitle"]  = RGBColor(*tpl["subtitle_color"]) if tpl.get("subtitle_color") is not None else "#ff914d"
            valid.append(tpl)
        else:
            warn("Template '" + tpl["id"] + "' : fichiers images manquants, ignore.")
    if not valid:
        warn("Aucun template valide trouve dans templates/"); sys.exit(1)
    return valid

def choose_template(templates):
    section_title("CHOIX DU TEMPLATE")
    for i, t in enumerate(templates, 1):
        print("    " + str(i) + ")  " + t["name"])
    print()
    while True:
        c = ask("Votre choix", "1")
        if c.isdigit() and 1 <= int(c) <= len(templates):
            return templates[int(c) - 1]
        warn("Entrez un numero entre 1 et " + str(len(templates)))


# =============================================================
#  API CLAUDE
# =============================================================

def call_claude(prompt, max_tokens=500):
    """Appel API Anthropic. Retourne uniquement le texte."""
    if not API_KEY:
        return "[Paragraphe IA indisponible : definissez ANTHROPIC_API_KEY]"
    body = json.dumps({
        "model": MODEL,
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}]
    }).encode("utf-8")
    req = urllib.request.Request(
        API_URL, data=body, method="POST",
        headers={
            "Content-Type":      "application/json",
            "x-api-key":         API_KEY,
            "anthropic-version": "2023-06-01",
        }
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            res = json.loads(r.read().decode("utf-8"))
            # Extraire le texte du premier bloc de type "text"
            for block in res.get("content", []):
                if isinstance(block, dict) and block.get("type") == "text":
                    return block["text"].strip()
            return "[Reponse IA vide]"
    except urllib.error.HTTPError as e:
        err = e.read().decode("utf-8", errors="replace")
        return "[Erreur API " + str(e.code) + ": " + err[:120] + "]"
    except Exception as e:
        return "[Erreur reseau: " + str(e) + "]"

def ai_step(name, has_img, title):
    progress("IA => " + name + " ...")
    note = "Une capture d'ecran illustre cette etape." if has_img else ""
    r = call_claude(
        "Redige un paragraphe professionnel de 4 a 6 phrases en francais "
        "pour la section intitulee '" + name + "' "
        "dans un rapport technique intitule '" + title + "'. "
        + note + " Style formel et precis, sans titre, sans liste a puces.",
        400)
    done_ok(); return r

def ai_intro(title, author, sup, steps):
    progress("IA => Introduction ...")
    ss = ", ".join(steps) if steps else "plusieurs etapes techniques"
    r  = call_claude(
        "Redige une introduction professionnelle en francais de 3 a 5 phrases "
        "pour un rapport intitule '" + title + "'. "
        "Auteur : " + author + ". Encadrant : " + sup + ". "
        "Etapes abordees : " + ss + ". Style formel, sans titre.",
        400)
    done_ok(); return r

def ai_conclusion(title, steps):
    progress("IA => Conclusion ...")
    ss = ", ".join(steps) if steps else "les etapes du rapport"
    r  = call_claude(
        "Redige une conclusion professionnelle en francais de 4 a 6 phrases "
        "pour un rapport intitule '" + title + "'. "
        "Etapes couvertes : " + ss + ". "
        "Synthetise les resultats, souligne les points cles, ouvre sur des perspectives. "
        "Sans titre.",
        500)
    done_ok(); return r


# =============================================================
#  WATERMARK (image de fond pleine page dans le header)
# =============================================================

def _inline_to_anchor_behind(inline_el, page_w_emu, page_h_emu):
    """
    Convertit un element wp:inline (image inline) en wp:anchor positionne
    en absolu sur la page avec behindDoc='1' (derriere le texte).
    """
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT",          "0")
    anchor.set("distB",          "0")
    anchor.set("distL",          "0")
    anchor.set("distR",          "0")
    anchor.set("simplePos",      "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc",      "1")   # <- DERRIERE le texte
    anchor.set("locked",         "0")
    anchor.set("layoutInCell",   "1")
    anchor.set("allowOverlap",   "0")

    sp = OxmlElement("wp:simplePos"); sp.set("x","0"); sp.set("y","0")
    anchor.append(sp)

    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom","page")
    po = OxmlElement("wp:posOffset"); po.text = "0"
    ph.append(po); anchor.append(ph)

    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom","page")
    po2 = OxmlElement("wp:posOffset"); po2.text = "0"
    pv.append(po2); anchor.append(pv)

    # Copier extent depuis inline (dimensions de l'image) mais forcer page size
    ext = OxmlElement("wp:extent")
    ext.set("cx", str(int(page_w_emu)))
    ext.set("cy", str(int(page_h_emu)))
    anchor.append(ext)

    ee = OxmlElement("wp:effectExtent")
    ee.set("l","0"); ee.set("t","0"); ee.set("r","0"); ee.set("b","0")
    anchor.append(ee)

    anchor.append(OxmlElement("wp:wrapNone"))

    # Copier docPr, cNvGraphicFramePr, graphic depuis l'inline original
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    for tag in ("docPr", "cNvGraphicFramePr"):
        el = inline_el.find("{%s}%s" % (NS_WP, tag))
        if el is not None:
            anchor.append(copy.deepcopy(el))

    NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    graphic = inline_el.find("{%s}graphic" % NS_A)
    if graphic is not None:
        anchor.append(copy.deepcopy(graphic))

    return anchor


def add_watermark_to_section(doc, section, img_path):
    """
    Ajoute une image pleine page en filigrane (derriere le texte)
    dans le header de la section donnee.
    """
    page_w = float(section.page_width)
    page_h = float(section.page_height)

    header = section.header
    header.is_linked_to_previous = False

    # Vider le header existant
    for p in list(header.paragraphs):
        for r in list(p.runs):
            r.clear()
        # Supprimer les drawings existants
        for drawing in p._p.findall(".//" + qn("w:drawing")):
            drawing.getparent().remove(drawing)

    # Paragraphe du header
    hpar = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hpar.paragraph_format.space_before = Pt(0)
    hpar.paragraph_format.space_after  = Pt(0)

    # Inserer l'image en inline (python-docx fait ca nativement)
    run = hpar.add_run()
    run.add_picture(str(img_path), width=int(page_w), height=int(page_h))

    # Trouver le drawing cree et convertir inline -> anchor behind
    drawing = hpar._p.find(".//" + qn("w:drawing"))
    if drawing is None:
        return

    NS_WP   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    inline  = drawing.find("{%s}inline" % NS_WP)
    if inline is not None:
        anchor = _inline_to_anchor_behind(inline, page_w, page_h)
        drawing.remove(inline)
        drawing.append(anchor)


# =============================================================
#  HELPERS DOCX
# =============================================================

def set_margins(sec, top_cm=2.5, bot_cm=2.0, left_cm=2.5, right_cm=2.0):
    sec.top_margin    = Cm(top_cm)
    sec.bottom_margin = Cm(bot_cm)
    sec.left_margin   = Cm(left_cm)
    sec.right_margin  = Cm(right_cm)

def add_paragraph_text(doc, text, size=11, bold=False, italic=False,
                       color=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before=4, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_thin_line(doc, color_hex):
    """Trait separateur horizontal."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    p.add_run().font.size = Pt(1)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot); pPr.append(pBdr)


# =============================================================
#  CONSTRUCTION DES PAGES
# =============================================================

def build_cover_pages(doc, tpl, title, author, supervisor, date_str):
    """
    Page unique de couverture : texte + filigrane intro en fond.
    (La page image pure a ete supprimee)
    """
    accent_hex = "{:02X}{:02X}{:02X}".format(*tpl["accent_color"])

    # ── Page de couverture unique ─────────────────────────────
    sec2 = doc.sections[0]
    sec2.page_width    = Inches(8.27)
    sec2.page_height   = Inches(11.69)
    set_margins(sec2, top_cm=0.5, bot_cm=2.0, left_cm=2.5, right_cm=2.0)
    sec2.header_distance = Emu(0)

    add_watermark_to_section(doc, sec2, tpl["_intro"])

    # Espace pour pousser le texte vers le bas (zone sans cercles)
    for _ in range(22):
        p = doc.add_paragraph("")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

    # Titre en gras
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(2)
    rt = p_title.add_run(title.upper())
    rt.font.size      = Pt(tpl.get("title_font_size", FONT_SIZE_COVER))
    rt.font.bold      = True
    rt.font.color.rgb = tpl["_accent"]

    # REPORT
    p_rep = doc.add_paragraph()
    p_rep.paragraph_format.space_before = Pt(0)
    p_rep.paragraph_format.space_after  = Pt(2)
    rr = p_rep.add_run("REPORT")
    rr.font.size      = Pt(tpl.get("report_font_size", FONT_SIZE_REPORT))
    rr.font.bold      = False
    rr.font.color.rgb = tpl["_secondary"]

    add_thin_line(doc, accent_hex)

    # Prepared By
    p_auth = doc.add_paragraph()
    p_auth.paragraph_format.space_before = Pt(4)
    p_auth.paragraph_format.space_after  = Pt(2)
    la = p_auth.add_run("Prepared By :\n")
    la.font.size = Pt(tpl.get("author_font_size", 21)); la.font.color.rgb = tpl["_secondary"]
    ra = p_auth.add_run(author)
    ra.font.size = Pt(tpl.get("author_name_font_size", 24)); ra.font.bold = True; ra.font.color.rgb = tpl["_dark"]

    # Supervised By
    p_sup = doc.add_paragraph()
    p_sup.paragraph_format.space_before = Pt(3)
    p_sup.paragraph_format.space_after  = Pt(2)
    ls = p_sup.add_run("Supervised By :\n")
    ls.font.size = Pt(tpl.get("supervisor_font_size", 21)); ls.font.color.rgb = tpl["_secondary"]
    rs = p_sup.add_run(supervisor)
    rs.font.size = Pt(tpl.get("supervisor_name_font_size", 24)); rs.font.bold = True; rs.font.color.rgb = tpl["_dark"]

    # Date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(6)
    rd = p_date.add_run(date_str)
    rd.font.size = Pt(tpl.get("date_font_size", 15)); rd.font.color.rgb = tpl["_secondary"]


    rt.font.name = FONT_NAME         # Titre principal
    rr.font.name = FONT_NAME         # mot REPORT
    la.font.name = FONT_NAME         # label "Prepared By :"
    ra.font.name = FONT_NAME         # nom auteur
    ls.font.name = FONT_NAME         # label "Supervised By :"
    rs.font.name = FONT_NAME         # nom superviseur
    rd.font.name = FONT_NAME         # date

    doc.add_page_break()


def _new_content_section(doc, tpl):
    """Cree une nouvelle section avec watermark page_bg."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width    = Inches(8.27)
    sec.page_height   = Inches(11.69)
    set_margins(sec, top_cm=2.5, bot_cm=2.0, left_cm=2.5, right_cm=2.0)
    sec.header_distance = Emu(0)
    add_watermark_to_section(doc, sec, tpl["_page"])
    return sec

def _write_subtitle(doc, tpl, text):
    """Ecrit un sous-titre (ex: nom d'une etape) dans le doc."""
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(6)
    p_sub.paragraph_format.space_after  = Pt(2)
    rs = p_sub.add_run(text)
    rs.font.size      = Pt(tpl.get("subtitle_font_size", FONT_SIZE_SUBTITLE))
    rs.font.bold      = True
    rs.font.underline   = True
    rs.font.color.rgb = tpl["_subtitle"]
    rs.font.name      = FONT_NAME

def _write_bloc(doc, tpl, subtitle, paragraph, images, fig_base, fig_counter):
    """Ecrit un bloc (subtitle + paragraphe + images) dans le doc. Retourne fig_counter mis a jour."""
    if subtitle:
        _write_subtitle(doc, tpl, subtitle)
    if paragraph:
        p_body = doc.add_paragraph(str(paragraph))
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_body.paragraph_format.space_before = Pt(4)
        p_body.paragraph_format.space_after  = Pt(8)
        for run in p_body.runs:
            run.font.size      = Pt(tpl.get("body_font_size", FONT_SIZE_BODY))
            run.font.color.rgb = tpl["_dark"]
            run.font.name      = FONT_NAME
    for img_path in images:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after  = Pt(3)
        try:
            p_img.add_run().add_picture(str(img_path), width=Inches(5.0))
        except Exception as e:
            p_img.add_run("[Image non chargee : " + str(e) + "]")
        fig_counter += 1
        cap = doc.add_paragraph("Figure " + str(fig_counter) + " - " + fig_base)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after  = Pt(8)
        for r in cap.runs:
            r.font.size      = Pt(tpl.get("caption_font_size", FONT_SIZE_CAPTION))
            r.font.italic    = True
            r.font.color.rgb = tpl["_secondary"]
            r.font.name      = FONT_NAME
    return fig_counter


def build_content_page(doc, tpl, heading, pages_data):
    """
    Construit une ou plusieurs pages pour une section.
    pages_data : liste de pages.
    Chaque page = liste de blocs = {"subtitle": str, "paragraph": str, "images": [Path, ...]}.
    Premiere page : titre + separateur + blocs.
    Pages suivantes : titre (suite) + blocs.
    """
    accent_hex  = "{:02X}{:02X}{:02X}".format(*tpl["accent_color"])
    fig_counter = 0
    for page_idx, blocs in enumerate(pages_data):
        _new_content_section(doc, tpl)
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(0)
        p_h.paragraph_format.space_after  = Pt(2)
        label = heading if page_idx == 0 else heading + "  (suite)"
        rh = p_h.add_run(label)
        rh.font.size      = Pt(tpl.get("heading_font_size", FONT_SIZE_TITLE))
        rh.font.bold      = True
        rh.font.color.rgb = tpl["_accent"]
        rh.font.name      = FONT_NAME

        add_thin_line(doc, accent_hex)
        for bloc in blocs:
            fig_counter = _write_bloc(
                doc, tpl,
                bloc.get("subtitle", ""),
                bloc.get("paragraph", ""),
                bloc.get("images", []),
                heading, fig_counter
            )


def build_outro_page(doc, tpl, heading, pages_data):
    """Conclusion : meme logique que build_content_page mais avec watermark outro."""
    accent_hex  = "{:02X}{:02X}{:02X}".format(*tpl["accent_color"])
    fig_counter = 0
    for page_idx, blocs in enumerate(pages_data):
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        sec.page_width    = Inches(8.27)
        sec.page_height   = Inches(11.69)
        set_margins(sec, top_cm=2.5, bot_cm=2.0, left_cm=2.5, right_cm=2.0)
        sec.header_distance = Emu(0)
        add_watermark_to_section(doc, sec, tpl["_outro"])
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(0)
        p_h.paragraph_format.space_after  = Pt(2)
        label = heading if page_idx == 0 else heading + "  (suite)"
        rh = p_h.add_run(label)
        rh.font.size      = Pt(tpl.get("heading_font_size", FONT_SIZE_TITLE))
        rh.font.bold      = True
        rh.font.color.rgb = tpl["_accent"]
        rh.font.name      = FONT_NAME

        add_thin_line(doc, accent_hex)
        for bloc in blocs:
            fig_counter = _write_bloc(
                doc, tpl,
                bloc.get("subtitle", ""),
                bloc.get("paragraph", ""),
                bloc.get("images", []),
                heading, fig_counter
            )




# =============================================================
#  ASSEMBLAGE FINAL
# =============================================================

def _text_to_pages_data(text):
    """Convertit un texte simple en pages_data a une seule page, un seul bloc."""
    return [[{"paragraph": text, "images": []}]]


def assemble(meta, tpl, intro, steps, conclusion):
    OUTPUT_DIR.mkdir(exist_ok=True)
    doc = Document()

    build_cover_pages(doc, tpl,
                      meta["title"], meta["author"],
                      meta["supervisor"], meta["date"])

    # intro peut etre str (ancien) ou pages_data (nouveau)
    intro_pd = intro if isinstance(intro, list) else _text_to_pages_data(intro)
    build_content_page(doc, tpl, "Introduction", intro_pd)

    for step in steps:
        # step["pages"] = pages_data si nouveau format
        pages_data = step.get("pages", _text_to_pages_data(step.get("paragraph", "")))
        build_content_page(doc, tpl, step["name"], pages_data)

    # conclusion peut etre str ou pages_data
    conclu_pd = conclusion if isinstance(conclusion, list) else _text_to_pages_data(conclusion)
    build_outro_page(doc, tpl, "Conclusion", conclu_pd)

    safe = "".join(c if (c.isalnum() or c in " _-") else "_" for c in meta["title"])
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    out  = OUTPUT_DIR / (safe + "_" + ts + ".docx")
    doc.save(str(out))
    return out


# =============================================================
#  COLLECTE INTERACTIVE
# =============================================================

def collect_metadata():
    section_title("INFORMATIONS DU RAPPORT")
    title = ask("Titre du rapport")
    while not title:
        warn("Le titre ne peut pas etre vide.")
        title = ask("Titre du rapport")
    author     = ask("Votre nom complet")
    supervisor = ask("Nom du superviseur / encadrant")
    date_str   = ask("Date", datetime.now().strftime("%d/%m/%Y"))
    return {"title": title, "author": author, "supervisor": supervisor, "date": date_str}


def collect_image(prompt="Chemin de l'image (png/jpg)"):
    """Demande un chemin d'image valide. Retourne Path ou None."""
    while True:
        raw = ask(prompt)
        if not raw:
            return None
        p = Path(raw.strip().strip('"').strip("'"))
        if p.exists():
            info("Image trouvee : " + p.name)
            return p
        warn("Fichier introuvable : " + str(p))
        if not ask_yesno("Reessayer ?", default=True):
            return None


def collect_bloc(bloc_num, step_name, report_title, page_num):
    """
    Collecte un bloc (paragraphe IA + images multiples) pour une page donnee.
    Retourne {"paragraph": str, "images": [Path, ...]}.
    """
    print()
    print("    --- Bloc " + str(bloc_num) + " ---")

    # sous-titre du bloc (optionnel)
    sub = "" 
    if ask_yesno("Ajouter un sous-titre pour ce bloc ?", default=False):
        sub = ask("Texte du sous-titre")
        if sub:
            print("    Sous-titre : " + sub)
            step_name = sub  # pour le prompt IA du paragraphe

    # Images du bloc
    images = []
    if ask_yesno("Ajouter une image pour ce bloc ?", default=False):
        p = collect_image()
        if p:
            images.append(p)
        while ask_yesno("Ajouter une autre image ?", default=False):
            p = collect_image()
            if p:
                images.append(p)

    # Paragraphe IA
    paragraph = ai_step(step_name, len(images) > 0, report_title)
    print()
    wrap_print(paragraph)
    print()
    if not ask_yesno("Conserver ce paragraphe ?", default=True):
        paragraph = collect_manual("votre paragraphe")

    return {"subtitle": sub, "paragraph": paragraph, "images": images}


def collect_page_blocs(step_name, report_title, page_num):
    """
    Collecte tous les blocs d'une page.
    Retourne liste de blocs.
    """
    print()
    print("  === Page " + str(page_num) + " de l'etape : " + step_name + " ===")
    blocs = []
    bloc_num = 1

    # Premier bloc obligatoire
    blocs.append(collect_bloc(bloc_num, step_name, report_title, page_num))
    bloc_num += 1

    # Blocs supplementaires sur cette meme page
    while ask_yesno("Ajouter un autre bloc sur cette meme page ?", default=False):
        blocs.append(collect_bloc(bloc_num, step_name, report_title, page_num))
        bloc_num += 1

    return blocs


def collect_steps(report_title):
    section_title("ETAPES DU RAPPORT")
    info("Tapez le nom de l'etape puis Entree.")
    info("Chaque etape peut avoir plusieurs pages.")
    info("Chaque page peut avoir plusieurs blocs (paragraphe + images).")
    info("Tapez  fin  pour terminer les etapes.\n")

    steps, idx = [], 1
    while True:
        print()
        print("  +-- Etape " + str(idx) + " " + "-" * 42)
        name = ask("Nom de l'etape (ou 'fin')")
        if name.lower() in ("fin", "done", "exit", "q", "quit"):
            break
        if not name:
            warn("Nom vide, recommencez."); continue

        # Collecte des pages de cette etape
        pages_data = []
        page_num   = 1

        pages_data.append(collect_page_blocs(name, report_title, page_num))
        page_num += 1

        while ask_yesno("Continuer cette etape sur une nouvelle page ?", default=False):
            pages_data.append(collect_page_blocs(name, report_title, page_num))
            page_num += 1

        steps.append({"name": name, "pages": pages_data})
        info("Etape " + str(idx) + " : " + str(page_num - 1) + " page(s), "
             + str(sum(len(p) for p in pages_data)) + " bloc(s) au total.")
        idx += 1
    return steps


# =============================================================
#  MAIN
# =============================================================

def main():
    banner()

    if not API_KEY:
        warn("ANTHROPIC_API_KEY non definie.")
        warn("  PowerShell : $env:ANTHROPIC_API_KEY='sk-ant-...'")
        warn("  Les paragraphes seront des placeholders.\n")

    # Chargement des templates
    templates = load_templates()
    tpl       = choose_template(templates)
    info("Template selectionne : " + tpl["name"])

    # Metadonnees
    meta = collect_metadata()

    # Introduction
    section_title("INTRODUCTION")
    if ask_yesno("Ecrire l'introduction manuellement ?", default=False):
        intro = collect_manual("l'introduction")
    else:
        intro = None   # generee apres collecte des etapes

    # Etapes
    steps      = collect_steps(meta["title"])
    step_names = [s["name"] for s in steps]

    # Intro IA (si pas manuelle)
    if intro is None:
        intro = ai_intro(meta["title"], meta["author"], meta["supervisor"], step_names)

    # Conclusion
    section_title("CONCLUSION")
    if ask_yesno("Ecrire la conclusion manuellement ?", default=False):
        conclusion = collect_manual("la conclusion")
    else:
        conclusion = ai_conclusion(meta["title"], step_names)

    # Generation du fichier
    section_title("GENERATION")
    progress("Assemblage du document Word ...")
    try:
        out = assemble(meta, tpl, intro, steps, conclusion)
        done_ok()
        print()
        print("=" * 58)
        print("  RAPPORT GENERE AVEC SUCCES !")
        print("  Fichier  : " + str(out))
        print("  Template : " + tpl["name"])
        print("  Etapes   : " + str(len(steps)))
        print("=" * 58)
        print()
    except Exception as e:
        import traceback
        print(" ECHEC")
        print()
        traceback.print_exc()


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n  Annule par l'utilisateur.\n")
        sys.exit(0)